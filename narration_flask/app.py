from flask import Flask, request, render_template, send_file
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Helper functions for text processing and document creation (same as provided in the original code)

def half_to_full_width(text):
    trans_table = str.maketrans(
        '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.',
        '０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ．'
    )
    return text.translate(trans_table)

def process_text(content, replacements):
    for pattern, replacement in replacements:
        content = re.sub(pattern, replacement, content, flags=re.MULTILINE)
    return content

def remove_first_duplicate_line(content):
    lines = content.splitlines()
    pattern = r'\b\d{4}\b'
    all_numbers = re.findall(pattern, content)
    count_numbers = {num: all_numbers.count(num) for num in set(all_numbers)}
    first_occurrences = {num: False for num in count_numbers if count_numbers[num] > 1}
    new_lines = []
    for line in lines:
        numbers_in_line = re.findall(pattern, line)
        if any(num in first_occurrences and not first_occurrences[num] for num in numbers_in_line):
            for num in numbers_in_line:
                if num in first_occurrences and not first_occurrences[num]:
                    first_occurrences[num] = True
                    break
            continue
        new_lines.append(line)
    return "\n".join(new_lines)

def normalize_blank_lines(content):
    new_lines = []
    previous_line_was_blank = False
    for line in content.split('\n'):
        if line.strip():
            new_lines.append(line)
            previous_line_was_blank = False
        elif not previous_line_was_blank:
            new_lines.append(line)
            previous_line_was_blank = True
    return "\n".join(new_lines)

def create_document(content, output_path, template_path):
    replacements = [
        (r'V\d+, \d+\n{2,}', r''),
        (r'(\d{2});(\d{2});(\d{2});(\d{2})', r'\2\3'),
        (r'(^(?:Ｎ|N|N)[\s　]+)(?=.+\n)', r''),
        (r'^\s+(?=(?:.+\n)[\s　])', r''),
        (r'^[\s　]+(?=\S+\n+)', r''),
        (r'(\d{4})\s-\s(\d{4})\n(V\d{1,2},\s\d)\n((?:.+(?:\n|))*)', r'\1　　N　　\4\n\n\2　　ON\n'),
        (r'(^(?!.*\d{4}(?: |　)*(?:N|ON)(?: |　)*.*).+$)', r'　　　　　　　　　\1'),
    ]
    content = process_text(content, replacements)
    content = remove_first_duplicate_line(content)
    content = normalize_blank_lines(content)
    content = half_to_full_width(content)

    doc = Document(template_path)
    doc.add_paragraph(content)

    if not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    regex = re.compile(r'[０-９]{4}　　ＯＮ')
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        if regex.search(original_text):
            paragraph.clear()
            last_end = 0
            for match in regex.finditer(original_text):
                paragraph.add_run(original_text[last_end:match.start()])
                highlighted_run = paragraph.add_run(match.group())
                highlighted_run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
                last_end = match.end()
            paragraph.add_run(original_text[last_end:])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Hiragino Maru Gothic Pro'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Maru Gothic Pro')
            run.font.size = Pt(10.5)

    doc.save(output_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'input_file' not in request.files or 'output_file' not in request.form:
            return 'No file part', 400
        
        input_file = request.files['input_file']
        output_file = request.form['output_file']
        if input_file.filename == '':
            return 'No selected file', 400
        
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_file.filename)
        input_file.save(input_path)
        
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{output_file}.docx")
        
        # テンプレートファイルのパスが正しいことを確認
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp.docx')
        
        with open(input_path, "r", encoding="utf-8") as file:
            content = file.read()
        
        create_document(content, output_path, template_path)
        
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        # コンソールにエラーを表示し、500エラーを返す
        print(f"Error: {e}")
        return 'Internal Server Error', 500

if __name__ == '__main__':
    app.run(debug=True)

